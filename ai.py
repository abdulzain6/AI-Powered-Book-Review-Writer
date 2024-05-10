from langchain.document_loaders import UnstructuredFileLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chat_models.openai import ChatOpenAI
from langchain.schema import Document
from langchain.chains.summarize import load_summarize_chain
from random import randrange
from prompts import REFINE_PROMPT, PROMPT
from docx import Document


class ReviewWriter:
    def __init__(self, openai_api_key: str) -> None:
        self.openai_api_key = openai_api_key
        
    def load_documents(self, file_path: str, chunk_size = 1500) -> list[Document]:
        docs = UnstructuredFileLoader(file_path=file_path).load()
        return RecursiveCharacterTextSplitter(chunk_size=chunk_size).split_documents(docs)

    def write_review(self, docs:list[Document], temperature: int, words_count: int, rules: str):
        llm = ChatOpenAI(openai_api_key=self.openai_api_key, temperature=temperature, request_timeout=60)
        chain = load_summarize_chain(llm=llm, chain_type="refine", verbose=True, refine_prompt=REFINE_PROMPT, question_prompt=PROMPT)
        return chain({"input_documents" : docs, "words_count" : words_count, "rules" : rules})["output_text"]
    
    def create_review_document(self, text: str, output_path: str):
        document = Document()
        heading = document.add_paragraph("Review:")
        heading.style = 'Heading1'
        document.add_paragraph(text)
        document.save(output_path)
        
    def run(self, infile: str, outfile: str, temperature: int, chunk_size: int, words_count: int, rules: str):
        docs = self.load_documents(infile, chunk_size)
        review = self.write_review(docs, temperature, words_count, rules)
        self.create_review_document(review, outfile)
        return review
